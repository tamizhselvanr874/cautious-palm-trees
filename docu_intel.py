# Final Version (Included Backoff and Jitter & Cleaned Code) - Final

import os
import re
import streamlit as st
import fitz  # PyMuPDF for PDF processing
from io import BytesIO
from docx import Document
from docx.shared import Pt  # type: ignore
import cv2
import numpy as np
from PIL import Image
import requests
import base64
import json
from azure.storage.blob import BlobServiceClient # type: ignore
import tempfile
import logging
import time
import random
import ast

# Azure OpenAI credentials
azure_endpoint = "https://theswedes.openai.azure.com/"
api_key = "783973291a7c4a74a1120133309860c0"
api_version = "2024-02-01"
model = "GPT-4o-mini"

logging.basicConfig(
    level=logging.INFO,  # Set log level to INFO
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler()],  # Send logs to the console
)

logging.basicConfig(
    level=logging.WARNING,  # Set log level to WARNING
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler()],  # Send logs to the console
)

logging.basicConfig(
    level=logging.CRITICAL,  # Set log level to CRITICAL
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler()],  # Send logs to the console
)

logging.basicConfig(
    level=logging.ERROR,  # Set log level to ERROR
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler()],  # Send logs to the console
)


# Azure Blob Storage credentials
connection_string = "DefaultEndpointsProtocol=https;AccountName=patentpptapp;AccountKey=4988gBY4D2RU4zdy1NCUoORdCRYvoOziWSHK9rOVHxy9pFXfKenRqyE/P+tpFpfmNObUm/zOCjeY+AStiCS3uw==;EndpointSuffix=core.windows.net"
container_name = "ppt-storage"

blob_service_client = BlobServiceClient.from_connection_string(connection_string)
# URL of your Azure function endpoint
azure_function_url = "https://doc2pdf.azurewebsites.net/api/HttpTrigger1"


# Function to convert PPT to PDF using Azure Function
def ppt_to_pdf(ppt_file, pdf_file):
    logging.info("Function to convert PPT to PDF using Azure Function")
    mime_type = (
        "application/vnd.openxmlformats-officedocument.presentationml.presentation"
    )
    headers = {
        "Content-Type": "application/octet-stream",
        "Content-Type-Actual": mime_type,
    }
    with open(ppt_file, "rb") as file:
        response = requests.post(azure_function_url, data=file.read(), headers=headers)
        if response.status_code == 200:
            with open(pdf_file, "wb") as pdf_out:
                pdf_out.write(response.content)
            return True
        else:
            st.error(f"File conversion failed with status code: {response.status_code}")
            logging.error(
                f"File conversion failed with status code: {response.status_code}"
            )
            st.error(f"Response: {response.text}")
            return False


patent_profanity_words = [
    "absolute",
    "absolutely",
    "all",
    "always",
    "authoritative",
    "authoritatively",
    "best",
    "biggest",
    "black hat",
    "black list",
    "blackhat",
    "blacklist",
    "broadest",
    "certain",
    "certainly",
    "chinese wall",
    "compel",
    "compelled",
    "compelling",
    "compulsorily",
    "compulsory",
    "conclusive",
    "conclusively",
    "constantly",
    "critical",
    "critically",
    "crucial",
    "crucially",
    "decisive",
    "decisively",
    "definitely",
    "definitive",
    "definitively",
    "determinative",
    "each",
    "earliest",
    "easiest",
    "embodiment",
    "embodiments",
    "entire",
    "entirely",
    "entirety",
    "essential",
    "essentially",
    "essentials",
    "every",
    "everything",
    "everywhere",
    "exactly",
    "exclusive",
    "exclusively",
    "exemplary",
    "exhaustive",
    "farthest",
    "finest",
    "foremost",
    "forever",
    "fundamental",
    "furthest",
    "greatest",
    "highest",
    "imperative",
    "imperatively",
    "important",
    "importantly",
    "indispensable",
    "indispensably",
    "inescapable",
    "inescapably",
    "inevitable",
    "inevitably",
    "inextricable",
    "inextricably",
    "inherent",
    "inherently",
    "instrumental",
    "instrumentally",
    "integral",
    "integrally",
    "intrinsic",
    "intrinsically",
    "invaluable",
    "invaluably",
    "invariably",
    "invention",
    "inventions",
    "irreplaceable",
    "irreplaceably",
    "key",
    "largest",
    "latest",
    "least",
    "littlest",
    "longest",
    "lowest",
    "major",
    "man hours",
    "mandate",
    "mandated",
    "mandatorily",
    "mandatory",
    "master",
    "maximize",
    "maximum",
    "minimize",
    "minimum",
    "most",
    "must",
    "nearest",
    "necessarily",
    "necessary",
    "necessitate",
    "necessitated",
    "necessitates",
    "necessity",
    "necessitating",
    "need",
    "needed",
    "needs",
    "never",
    "newest",
    "nothing",
    "nowhere",
    "obvious",
    "obviously",
    "oldest",
    "only",
    "optimal",
    "ought",
    "overarching",
    "paramount",
    "perfect",
    "perfected",
    "perfectly",
    "perpetual",
    "perpetually",
    "pivotal",
    "pivotally",
    "poorest",
    "preferred",
    "purest",
    "required",
    "requirement",
    "requires",
    "requisites",
    "shall",
    "shortest",
    "should",
    "simplest",
    "slaves",
    "slightest",
    "smallest",
    "tribal knowledge",
    "ultimate",
    "ultimately",
    "unavoidable",
    "unavoidably",
    "unique",
    "uniquely",
    "unrivalled",
    "urgent",
    "urgently",
    "valuable",
    "very",
    "vital",
    "vitally",
    "white hat",
    "white list",
    "whitehat",
    "whitelist",
    "widest",
    "worst",
]


def encode_image(image):
    buffered = cv2.imencode(".jpg", image)[1]
    return base64.b64encode(buffered).decode("utf-8")


def extract_titles_from_images(image_content):
    slide_data = []
    logging.info("Function to extract titles from images")
    headers = {"Content-Type": "application/json", "api-key": api_key}

    for image_data in image_content:
        slide_number = image_data["slide_number"]
        base64_image = encode_image(image_data["image"])

        # Use LLM to generate a title for the slide based on the image
        prompt = "What is the title of the slide based on the given image?"
        data = {
            "model": model,
            "messages": [
                {
                    "role": "system",
                    "content": "You are a slide titles extraction model [Note: Only return the slide Title without any additional generated text]",
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{base64_image}"
                            },
                        },
                    ],
                },
            ],
            "max_tokens": 100,
            "temperature": 0.3,
        }

        response = requests.post(
            f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version={api_version}",
            headers=headers,
            data=json.dumps(data),
        )

        if response.status_code == 200:
            slide_title = response.json()["choices"][0]["message"]["content"]
        else:
            slide_title = "Untitled Slide"
            logging.warning("No Slide Title has been Found!")

        slide_data.append(
            {
                "slide_number": slide_number,
                "title": slide_title.strip(),
                "image": image_data["image"],
            }
        )

    return slide_data


# Function to generate insights for images via LLM
def generate_image_insights(
    image_content,
    text_length,
    low_quality_slides,
    system_prompt,
    slide_data,
):
    insights = []
    logging.info("Function to generate Image Insights")
    # Set temperature based on text length
    temperature = (
        0.0 if text_length == "Standard" else 0.5 if text_length == "Blend" else 0.7
    )

    for image_data in image_content:
        slide_number = image_data["slide_number"]
        if slide_number in low_quality_slides:
            continue  # Skip low-quality slides

        base64_image = encode_image(image_data["image"])

        # Determine how many images are on the slide
        images_on_slide = [
            img for img in image_content if img["slide_number"] == slide_number
        ]
        image_ref = f"figure {slide_number}"
        if len(images_on_slide) > 1:
            image_ref += f"({chr(97 + images_on_slide.index(image_data))})"

        # Get the slide title for mapping
        slide_title = next(
            (
                slide["title"]
                for slide in slide_data
                if slide["slide_number"] == slide_number
            ),
            "Untitled Slide",
        )

        headers = {
            "Content-Type": "application/json",
            "api-key": api_key,
            "Cache-Control": "no-cache",
            "Pragma": "no-cache",
        }

        prompt = f"""
                    Important Note: Avoid using words like 'contain', 'necessity', 'necessary', 'necessitate', 'contain' , 'contains' , 'consist,' 'explore' , 'key component' , 'revolutionizing',  'innovative' , or similar terms. Use alternatives instead. Return the content in one paragraph only.
                    Important Note: Avoid expanding abbreviations unless instructed in the given slide. Only expand abbreviations once.

                    Figure Detection:
                    ~~/ Identify and list all figures (diagrams, sketches, flowcharts) on the slide. Each figure, even if stacked or side by side, should be treated as separate.
                        If the figure on the slide has multiple individual parts that are connected as one overall figure, treat all these parts as a single figure and reference it using the slide number.
                        Note: If the figure number is already mentioned on the slide, ignore it and instead use "{slide_number}" as the figure number.
                        
                        If only figures are present, follow these steps based on the slide title:
                            Mention and reference each figure in order (e.g., "Referring to Figure {slide_number}(a), Figure {slide_number}(b)..."). Clearly explain each figure, covering its role and relevance.
                            Note: If the figure number is already mentioned on the slide, ignore it and instead use "{slide_number}" as the figure number.
                            Steps:
                        ```/ Reference figures in order: "Referring to Figure {slide_number}(a), Figure {slide_number}(b)…" Each figure must be mentioned individually.
                                Check that figures are referenced in order before any detailed descriptions.
                                Flag any issues if figures are missing or combined improperly (e.g., "the figures" instead of individual references).
                                For complex or overlapping figures, explain their relationships clearly, such as "Figure {slide_number}(a) interacts with Figure {slide_number}(b)..."
                                After Figure Reference follow these instructions based on the slide title:
                                    If the title includes "Background" or "Motivation": Start with "Referring to Figure {slide_number}, In this prior solutions include..." and focus only on prior solutions.
                                    If the title includes "Invention" or "Proposal": Start with "Referring to Figure {slide_number}, In this present disclosure includes..." and focus on the invention or proposal.
                                    If the title does not contain "Background", "Motivation", "Invention", "Proposal": Start with "Referring to Figure {slide_number}, In this aspect..." followed by a detailed explanation.
                                For Graphs:
                                    Describe the x and y axes and explain the overall meaning.
                                For Images:
                                    Identify angles, depth, and spatial relationships for images with perspective views. Refer to images specifically (avoid terms like "left" or "right" figure).
                                Natural Flow:
                                    Avoid phrases like "The slide shows..." or "The image presents..." to ensure a natural flow. /```

                                When generating content, avoid using the words "necessary" and "contain" and its related words. Instead, use alternatives like "required."
                                                                                
                    Style Guide:
                    ```/ Use passive voice, except for discussing the present disclosure (use active voice like "provides").
                        Replace "Million" and "Billion" with "1,000,000" and "1,000,000,000."
                        Avoid using "invention" or "objective," replace with "present disclosure."
                        Use technical terms and Avoid expanding abbreviations unless instructed. Only expand abbreviations once.
                        Turn bullet points into sentences without summarizing them. /``` ~~/ """

        data = {
            "model": model,
            "messages": [
                {
                    "role": "system",
                    "content": f""" {system_prompt}\n\n 
                                    Your task is to generate content based on the provided slide while adhering to the following instructions: 
                                    {prompt} """,
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{base64_image}"
                            },
                        }
                    ],
                },
            ],
            "temperature": temperature,
        }

        attempt = 0
        try:
            response = requests.post(
                f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version={api_version}",
                headers=headers,
                data=json.dumps(data),
            )

            llm_result = response.json()["choices"][0]["message"]["content"]
            res_content = replace_disallowed_words(llm_result)

            if response.status_code == 200:
                insights.append(
                    {
                        "slide_number": slide_number,
                        "slide_title": slide_title,
                        "insight": res_content,
                    }
                )
            else:
                insights.append(
                    {
                        "slide_number": slide_number,
                        "slide_title": slide_title,
                        "insight": "Error generating insight.",
                    }
                )
                logging.critical(
                    "Error generating image insight.... (Response status is Failed)"
                )
        except requests.exceptions.RequestException as e:
            attempt += 1
            if attempt < 4:
                logging.error(f"Error generating content': {e}")
                # Exponential backoff with random jitter
                backoff_time = (2**attempt) + random.uniform(0, 1)
                time.sleep(backoff_time)

    return insights


def continued_title_check(slide_data):
    continued = []
    logging.info("Function to continued title check")
    system_prompt = "You are tasked with identifying slides that share the same title. Once you detect identical titles across slides, check if any of these titles are followed by '(continued...)'. If a slide with the same title includes '(continued...)', return all the slides with the identical title, including those marked with '(continued...)' and those without. Ensure that slides are grouped appropriately based on title similarity and the presence of '(continued...)'. "

    t_value = []
    for sl in slide_data:
        slide_number = sl["slide_number"]
        title = sl["title"]
        t_value.append(f"{slide_number} : {title}")

    headers = {
        "Content-Type": "application/json",
        "api-key": api_key,
        "Cache-Control": "no-cache",
        "Pragma": "no-cache",
    }

    prompt = f""" 
    Check for slides with identical titles. If multiple slides share the same title, verify if any of these slides have the same title followed by '(Continued...)'. If any of the identical titled slides include '(Continued...)', return all slides with that title, both with and without '(Continued...)' in the title.        
    Note: To qualify as identical titles, at least one of the titles must include "Continued...". If "Continued..." is not present in any of the titles, they should not be considered identical, and the response should be: "No".
    Slide Data: {t_value}
    Example output format for identical titled slides Matched [Only return Output like this]: Yes, [3,4,5]
    Example output format for No identical titled slides Matched [Only return Output like this]: No
    Example output format for more than one identical titled slides Matched [Only return Output like this]: Yes, [3,4,5],[6, 7, 8],[9, 11, 12, 13],...
    """

    data = {
        "model": model,
        "messages": [
            {"role": "system", "content": f"{system_prompt}"},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0,
        "max_tokens": 100,
    }

    try:
        response = requests.post(
            f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version={api_version}",
            headers=headers,
            json=data,
        )

        response_data = response.json()  # Convert the response to JSON

        # Check if 'choices' key is in the response
        if "choices" in response_data:
            sp = response_data["choices"][0]["message"][
                "content"
            ]  # Access the message content
            # st.sidebar.write(sp)
            sp = sp.split("Yes,", 1)[1].strip()
            #  st.sidebar.write(sp)

            if sp:  # Check if the response contains "Yes"
                continued.append({"set_of_slides": sp})
        else:
            logging.error("Error: 'choices' key not found in the response.")
            st.write("Error: 'choices' key not found in the response.")

    except Exception as e:
        # Developer's error handling in console
        logging.warning(f"Error: {str(e)}")

    return continued


def generate_text_insights(
    text_content, text_length, low_quality_slides, slide_data, system_prompt
):
    logging.info("Function to generate Text Insights")
    headers = {
        "Content-Type": "application/json",
        "api-key": api_key,
        "Cache-Control": "no-cache",
        "Pragma": "no-cache",
    }
    insights = []
    base_delay = 1
    max_delay = 32

    # Set temperature based on text_length
    if text_length == "Standard":
        temperature = 0.0
    elif text_length == "Blend":
        temperature = 0.5
    elif text_length == "Creative":
        temperature = 0.7

    for slide in text_content:
        slide_number = slide["slide_number"]
        if slide_number in low_quality_slides:
            continue  # Skip low-quality slides

        slide_title = next(
            (
                slide["title"]
                for slide in slide_data
                if slide["slide_number"] == slide_number
            ),
            "Untitled Slide",
        )
        slide_text = slide["text"]

        prompt = f"""{system_prompt}\n
        \\\\
        Important Note: Avoid using words like 'contain', 'necessity,' 'necessary,' 'necessitate,' 'contain' , 'contains' , 'consist,' 'explore,' 'key component,' 'revolutionizing,' 'innovative,' or similar terms. Use alternatives instead. Return the content in one paragraph only.
        Important Note: Avoid expanding abbreviations unless instructed in the given slide. Only expand abbreviations once.

        I want you to begin with one of the following phrases based on the slide title: 
        
        (a) If the title includes "Invention" or "Proposal", start with:
        "The present disclosure includes..."
        Focus on the proposal or invention, without mentioning background information.

        (b) If the title includes "Background" or "Motivation," start with:
        "The prior solutions include..."
        Focus on prior solutions only, without including proposals.

        (c) If the title doesn't include "Background" or "Proposal," start with:
        "Aspects of the present disclosure include..."
        Focus on the slide's main points, without mentioning prior solutions or proposals. 
 
        The information should be delivered in a structured, clear, and concise paragraph while avoiding phrases like 'The slide presents,' 'discusses,' 'outlines,' or 'content.' Summarize all major points without bullet points.  
        When generating content, avoid using the words "necessary" and "contain" and its related words. Instead, use alternatives like "required."
        
        Note: Turn bullet points into sentences without summarizing them and make sure to mention all the reference numbers included in the points.
        
        Follow these detailed style guidelines for the generated content:          
            (a) Remove all listed profanity words: {patent_profanity_words}\n. 
            (b) Use passive voice, except for discussing the present disclosure (use active voice like "provides").
            (c) Replace "Million" and "Billion" with "1,000,000" and "1,000,000,000."
            (d) Avoid using "invention" or "objective," replace with "present disclosure."
            (e) Use detailed technical jargon.
            (f) Organize explanations systematically with terms like "defined as" or "for example."
            (g) Turn bullet points into sentences without summarizing them.
            (h) Maintain exact wording—don't replace terms with synonyms.
            (i) Use definitive language when discussing the current disclosure.
            (j) Ensure accurate representation of figures, flowcharts, and equations.
            (k) Avoid specific words like "revolutionizing" or "innovative."
            (l) Use technical terms and Avoid expanding abbreviations unless instructed. Only expand abbreviations once. 
            
        Important Note: Return content only in a single paragraph.
        Important Note: Give importance to equations that are presented in the Slide.
        Important Note: Don't consider equation as Images.
        Important Note: Do not expand abbreviations on its own unless mentioned in the slide. 
        Important Note: Only expand abbreviations one time throughout the entire content.\\\\
        
        Slide Text: ```{slide_text}```
        """

        if text_length == "Standard":
            prompt += "\n\nGenerate a short paragraph"
        elif text_length == "Blend":
            prompt += "\n\nGenerate a medium-length paragraph"
        elif text_length == "Creative":
            prompt += "\n\nGenerate a longer paragraph."

        data = {
            "model": model,
            "messages": [
                {"role": "system", "content": f"""{system_prompt}"""},
                {"role": "user", "content": prompt},
            ],
            "temperature": temperature,
        }

        try:
            response = requests.post(
                f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version={api_version}",
                headers=headers,
                json=data,
            )

            llm_result = response.json()["choices"][0]["message"]["content"]
            res_content = replace_disallowed_words(llm_result)

            if response.status_code == 200:
                insights.append(
                    {
                        "slide_number": slide_number,
                        "slide_title": slide_title,
                        "insight": res_content,
                    }
                )
            else:
                insights.append(
                    {
                        "slide_number": slide_number,
                        "slide_title": slide_title,
                        "insight": "Error generating insight.",
                    }
                )
                logging.critical(
                    "Error generating text insight.... (Response status is Failed)"
                )
        except requests.exceptions.RequestException as e:
            attempt = 2
            delay = min(max_delay, base_delay * (2**attempt))
            jitter = random.uniform(0, delay)
            logging.warning(
                f"Retrying in {jitter:.2f} seconds (attempt {attempt}) due to error: {e}"
            )
            time.sleep(jitter)

    return insights


def generate_prompt(overall_theme):
    logging.info("Function to generate Dynamic system prompt")
    headers = {
        "Content-Type": "application/json",
        "api-key": api_key,
        "Cache-Control": "no-cache",
        "Pragma": "no-cache",
    }

    # Generate an overall theme of the following document content: {text_content}
    prompt = f"Create a perfect system prompt based on the given content: {overall_theme}\n [Note: Return output in single line starting with 'You are a Patent Attorney specializing..]"

    data = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "You are a Patent Attorney specializing in generating content based on the document content",
            },
            {"role": "user", "content": prompt},
        ],
        "max_tokens": 600,
        "temperature": 0.3,
    }

    response = requests.post(
        f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version={api_version}",
        headers=headers,
        data=json.dumps(data),
    )

    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    else:
        logging.info("Default system prompt has been returned")
        return "You are a Patent Attorney specializing in generating content based on the document content"


# Function to detect images, flowcharts, and diagrams from the PDF focused to Title
def extract_slide_images_for_title_extraction(pdf_path):
    logging.info("Function to extract slide images for title extraction")

    doc = fitz.open(pdf_path)
    image_content = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap()
        img_data = pix.samples
        img = Image.frombytes("RGB", [pix.width, pix.height], img_data)
        img_np = np.array(img)

        gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
        thresh = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 11, 2
        )

        contours, _ = cv2.findContours(
            thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE
        )
        # significant_contours = [cnt for cnt in contours if cv2.contourArea(cnt) > 1000]

        image_content.append({"slide_number": page_num + 1, "image": img_np})

    return image_content


# Function to detect images, flowcharts, and diagrams from the PDF
def detect_images_from_pdf(pdf_path):
    logging.info("Function to detect images from PDF")
    doc = fitz.open(pdf_path)
    image_content = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap()
        img_data = pix.samples
        img = Image.frombytes("RGB", [pix.width, pix.height], img_data)
        img_np = np.array(img)

        gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
        thresh = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 11, 2
        )

        contours, _ = cv2.findContours(
            thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE
        )
        significant_contours = [cnt for cnt in contours if cv2.contourArea(cnt) > 1000]

        if len(significant_contours) > 0:
            image_content.append({"slide_number": page_num + 1, "image": img_np})
            # st.success(page_num + 1)

    return image_content


def extract_text_and_titles_from_pdf(pdf_path):
    logging.info("Function to extract text and titles from PDF using Fitz (PyMuPDF)")
    doc = fitz.open(pdf_path)
    slide_data = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        blocks = page.get_text("dict").get("blocks", [])
        page_text = ""
        slide_title = ""

        for block in blocks:
            if "lines" in block:  # Check if 'lines' key exists in the block
                for line in block["lines"]:
                    for span in line["spans"]:
                        text = span["text"].strip()
                        # Improved title detection: Check for larger font size, bold text, and exclude symbols
                        if (
                            (span["size"] > 14 or span["flags"] & 2)
                            and len(text) > 2
                            and not any(char in text for char in ["•", "-", "*"])
                        ):
                            if (
                                not slide_title
                            ):  # Only take the first valid occurrence as the title
                                slide_title = text
                        page_text += text + " "

        slide_data.append(
            {
                "page_number": page_num + 1,
                "title": slide_title if slide_title else "Untitled Slide",
                "content": page_text.strip(),
            }
        )

    return slide_data


def generate_continue_insights(
    text_length,
    continued_check,
    image_slides,
    text_content,
    slide_data,
    system_prompt,
):
    insights = []

    # Set temperature based on text length
    temperature = (
        0.0 if text_length == "Standard" else 0.5 if text_length == "Blend" else 0.7
    )

    # Process each continued_check dictionary entry for multiple slide sets
    for check_entry in continued_check:
        if isinstance(check_entry, dict):
            # Parse the "set_of_slides" string and ensure it is a list of lists
            continued_slide_sets = ast.literal_eval(
                check_entry.get("set_of_slides", "")
            )

            # If the result is a single list (e.g., [2, 3]), wrap it in another list
            if isinstance(continued_slide_sets[0], int):
                continued_slide_sets = [continued_slide_sets]

            # Process each slide set individually
            for continued_slide_numbers in continued_slide_sets:
                combined_title = None
                combined_images = []
                combined_text = []

                # Combine images and text for the current slide set
                for image_data in image_slides:
                    slide_number = image_data["slide_number"]
                    if slide_number in continued_slide_numbers:
                        slide_number_img = slide_number
                        combined_images.append(image_data)
                        if combined_title is None:
                            combined_title = (
                                next(
                                    (
                                        slide["title"]
                                        for slide in slide_data
                                        if slide["slide_number"] == slide_number_img
                                    ),
                                    "Untitled Slide",
                                )
                                .replace("(Continued)", "")
                                .strip()
                            )

                for text_data in text_content:
                    slide_number = text_data["slide_number"]
                    if slide_number in continued_slide_numbers:
                        combined_text.append(text_data)
                        if combined_title is None:
                            combined_title = (
                                next(
                                    (
                                        slide["title"]
                                        for slide in slide_data
                                        if slide["slide_number"] == slide_number
                                    ),
                                    "Untitled Slide",
                                )
                                .replace("(Continued)", "")
                                .strip()
                            )

                if combined_images and combined_title and combined_text:
                    base64_images = [
                        encode_image(img["image"]) for img in combined_images
                    ]
                    combined_slide_ref = ",".join(map(str, continued_slide_numbers))
                    logging.info("Function to generate Continued Image Insights")

                    # st.error(slide_number_img)
                    prompt = f"""{system_prompt}
                                                     
                            Objective:
                                Generate a detailed paragraph based on the provided slides, integrating both textual content and visual elements seamlessly. 
                                The response should prioritize the text content and use it explain the figures or images or diagrams or graphs or charts or equations. 
                                Ensure the explanations are well-structured and focused on integrating the textual information along with the figures or images or diagrams or graphs or charts or equations content.
                                     
                                Slide Information:
                                    Slide Reference: {slide_number_img}

                            Let's think step by step:

                            Please provide a comprehensive explanation of all figures (including diagrams, drawings, and sketches) mentioned in the following text:                        
                            Refer to the following text to provide a detailed explanation of the mentioned figures or images or diagrams or graphs or charts: `````{combined_text }`````\n\n
                            Your explanation should:
                            Include all text content: Ensure that every detail from the text is fully presented in your explanation.
                            Identify and describe all elements in the figures: Carefully recognize and explain each component in the figures, including any labels, legends, and annotations.
                            Focus on the operational flow: Provide a detailed explanation of the working process depicted in the figures, describing how each element interacts within the system.
                            Use exact terms from labels and legends: Maintain consistency by using the precise terms provided in the figures' labels or legends, without referencing phrases like "labelled parts" or "as indicated by the legend".
                            When generating content, avoid using the words "necessary" and "contain" and its related words. Instead, use alternatives like "required."
                            Supplement with additional details from the figures: Incorporate any extra information from the figures that isn't explicitly mentioned in the text.
                            Describe spatial relationships without directional terms: When explaining spatial aspects in the figures, focus on angles, depth, and interactions without using directional words like "left" or "right".
                            Explain roles and relevance: Explain the significance of each element and how it contributes to the overall process or concept being described.
                                                                            
                            Important Note: Avoid using words like 'contain', 'necessity', 'necessary', 'necessitate', 'contain' , 'contains' , 'consist,' 'explore' , 'key component' , 'revolutionizing',  'innovative' , or similar terms. Use alternatives instead. Return the content in one paragraph only.
                            Important Note: Avoid expanding abbreviations unless instructed in the given slide. Only expand abbreviations once. 

                                Figure Detection:
                                ~~/ Identify and list all figures (diagrams, sketches, flowcharts) on the slide. Each figure, even if stacked or side by side, should be treated as separate.
                                    If the figure on the slide has multiple individual parts that are connected as one overall figure, treat all these parts as a single figure and reference it using the slide number.
                                    Note: If the figure number is already mentioned on the slide, ignore it and instead use "{slide_number_img}" as the figure number.
                                    
                                Mention and reference each figure in order (e.g., "Referring to Figure {slide_number_img}(a), Figure {slide_number_img}(b)..."). Clearly explain each figure, covering its role and relevance.
                                Note: If the figure number is already mentioned on the slide, ignore it and instead use "{slide_number_img}" as the figure number.
                                Steps:
                            ```/ Reference figures in order: "Referring to Figure {slide_number_img}(a), Figure {slide_number_img}(b)…" Each figure must be mentioned individually.
                                    Check that figures are referenced in order before any detailed descriptions.
                                    After referencing, describe each figure individually: "Figure {slide_number_img}(a) illustrates...", "Figure {slide_number_img}(b) shows..." Explain the figures in detail.
                                    Flag any issues if figures are missing or combined improperly (e.g., "the figures" instead of individual references).
                                    For complex or overlapping figures, explain their relationships clearly, such as "Figure {slide_number_img}(a) interacts with Figure {slide_number_img}(b)..."  
                                    After Figure Reference follow these instructions based on the slide title from the Image slide:
                                    If Image the title includes 'Discussion': Start with "Referring to Figure {slide_number_img}, In this prior solutions include..." and focus only on prior solutions.
                                    If Image the title includes "Background" or "Motivation": Start with "Referring to Figure {slide_number_img}, In this prior solutions include..." and focus only on prior solutions.
                                    If Image the title includes "Invention" or "Proposal": Start with "Referring to Figure {slide_number_img}, In this present disclosure includes..." and focus on the invention or proposal.                         
                                    
                                    For Graphs:
                                        Describe the x and y axes and explain the overall meaning.
                                    For Images:
                                        Identify angles, depth, and spatial relationships for images with perspective views. Refer to images specifically (avoid terms like "left" or "right" figure).
                                    Natural Flow:
                                        Avoid phrases like "The slide shows..." or "The image presents..." to ensure a natural flow. /```
                                        

                                    
                        Style Guide:
                        ```/ Use passive voice, except for discussing the present disclosure (use active voice like "provides"). 
                            Replace "Million" and "Billion" with "1,000,000" and "1,000,000,000."
                            Avoid using "invention" or "objective," replace with "present disclosure."
                            Use technical terms and Avoid expanding abbreviations unless instructed. Only expand abbreviations once.
                            Turn bullet points into sentences without summarizing them. /```
                Slide:
                    """

                    # messages = []
                    for img_b64 in base64_images:
                        print(
                            "-----------------------------------------------------------------------------------------------------------------"
                        )

                        # messages.append({"role": "user", "content": {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}}})

                    data = {
                        "model": model,
                        "messages": [
                            {"role": "system", "content": f"""{prompt}"""},
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": f"data:image/png;base64,{img_b64}"
                                        },
                                    },
                                ],
                            },
                        ],
                        "temperature": temperature,
                    }

                    headers = {
                        "Content-Type": "application/json",
                        "api-key": api_key,
                        "Cache-Control": "no-cache",
                        "Pragma": "no-cache",
                    }

                    response = requests.post(
                        f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version={api_version}",
                        headers=headers,
                        json=data,
                    )

                    # response_data = response.json()
                    # sp = response_data['choices'][0]['message']['content']
                    # st.write(sp)

                    llm_result = response.json()["choices"][0]["message"]["content"]
                    res_content = replace_disallowed_words(llm_result)

                    # Process the response
                    if response.status_code == 200:
                        insights.append(
                            {
                                "slide_number": combined_slide_ref,
                                "slide_title": combined_title,
                                "insight": res_content,
                            }
                        )
                    else:
                        insights.append(
                            {
                                "slide_number": combined_slide_ref,
                                "slide_title": combined_title,
                                "insight": response.json(),
                            }
                        )
                        logging.critical(
                            "Error generating combined_image insight.... (Response status is Failed)"
                        )

                if combined_title and combined_text and combined_images == []:
                    logging.info("Function to generate Continued Text Insights")
                    combined_slide_ref = ",".join(map(str, continued_slide_numbers))
                    prompt = f"""{system_prompt}
                                                            
                                Objective:
                                                                                
                                Important Note: Avoid using words like 'contain', 'necessity', 'necessary', 'necessitate', 'contain' , 'contains' , 'consist,' 'explore' , 'key component' , 'revolutionizing',  'innovative' , or similar terms. Use alternatives instead. Return the content in one paragraph only.
                                Important Note: Avoid expanding abbreviations unless instructed in the given slide. Only expand abbreviations once. 

                                I want you to begin with one of the following phrases based on the slide title: 
                                
                                (a) If the title includes "Invention" or "Proposal", start with:
                                "The present disclosure includes..."
                                Focus on the proposal or invention, without mentioning background information.

                                (b) If the title includes "Background" or "Motivation," start with:
                                "The prior solutions include..."
                                Focus on prior solutions only, without including proposals.

                                (c) If the title doesn't include "Background" or "Proposal," start with:
                                "Aspects of the present disclosure include..."
                                Focus on the slide's main points, without mentioning prior solutions or proposals. 
                        
                                The information should be delivered in a structured, clear, and concise paragraph while avoiding phrases like 'The slide presents,' 'discusses,' 'outlines,' or 'content.' Summarize all major points without bullet points.  
                                When generating content, avoid using the words "necessary" and "contain" and its related words. Instead, use alternatives like "required."
                                
                                Note: Turn bullet points into sentences without summarizing them and make sure to mention all the reference numbers included in the points.
                                
                                Follow these detailed style guidelines for the generated content:          
                                    (a) Remove all listed profanity words: {patent_profanity_words}\n. 
                                    (b) Use passive voice, except for discussing the present disclosure (use active voice like "provides").
                                    (c) Replace "Million" and "Billion" with "1,000,000" and "1,000,000,000."
                                    (d) Avoid using "invention" or "objective," replace with "present disclosure."
                                    (e) Use detailed technical jargon.
                                    (f) Organize explanations systematically with terms like "defined as" or "for example."
                                    (g) Turn bullet points into sentences without summarizing them.
                                    (h) Maintain exact wording—don't replace terms with synonyms.
                                    (i) Use definitive language when discussing the current disclosure.
                                    (j) Avoid specific words like "revolutionizing" or "innovative."
                                    (k) Use technical terms and Avoid expanding abbreviations unless instructed. Only expand abbreviations once. 
                                
                                When generating content, avoid using the words "necessary" and "contain" and its related words. Instead, use alternatives like "required."
                                            
                                Style Guide:
                                ```/ Use passive voice, except for discussing the present disclosure (use active voice like "provides"). 
                                    Replace "Million" and "Billion" with "1,000,000" and "1,000,000,000."
                                    Avoid using "invention" or "objective," replace with "present disclosure."
                                    Use technical terms and Avoid expanding abbreviations unless instructed. Only expand abbreviations once.
                                    Turn bullet points into sentences without summarizing them. /```
                                
                                While generating  do not lose any of the given Slide Text, Include all the information from the given Text.
                                
                                Important Note: Return content only in a single paragraph.
                                Important Note: Give importance to all equations that are presented in the Slide.
                                Important Note: Don't consider equation as Images.
                                Important Note: Do not expand abbreviations on its own unless mentioned in the slide. 
                                Important Note: Only expand abbreviations one time throughout the entire content.\\\\
                
                        Slide Text: ```{combined_text}```
                """

                    data = {
                        "model": model,
                        "messages": [
                            {
                                "role": "system",
                                "content": f"""{system_prompt}
                                                            When generating content, avoid using the words "necessary" and "contain" and its related words. Instead, use alternatives like "required." """,
                            },
                            {"role": "user", "content": prompt},
                        ],
                        "temperature": temperature,
                    }

                    headers = {
                        "Content-Type": "application/json",
                        "api-key": api_key,
                        "Cache-Control": "no-cache",
                        "Pragma": "no-cache",
                    }

                    response = requests.post(
                        f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version={api_version}",
                        headers=headers,
                        json=data,
                    )

                    # response_data = response.json()
                    # sp = response_data['choices'][0]['message']['content']
                    # st.write(sp)

                    llm_result = response.json()["choices"][0]["message"]["content"]
                    res_content = replace_disallowed_words(llm_result)

                    # Process the response
                    if response.status_code == 200:
                        insights.append(
                            {
                                "slide_number": combined_slide_ref,
                                "slide_title": combined_title,
                                "insight": res_content,
                            }
                        )
                    else:
                        insights.append(
                            {
                                "slide_number": combined_slide_ref,
                                "slide_title": combined_title,
                                "insight": response.json(),
                            }
                        )
                        logging.critical(
                            "Error generating combined_text insight.... (Response status is Failed)"
                        )
    return insights


# Function to generate an overall theme based on extracted text
def generate_overall_theme(text_content):
    logging.info("Function to generate an overall theme based on extracted text")

    headers = {
        "Content-Type": "application/json",
        "api-key": api_key,
        "Cache-Control": "no-cache",
        "Pragma": "no-cache",
    }
    prompt = f"Analysis and identify the domain and subject of the patent/Invention and then generate an overall theme of the following document content: {text_content}"
    data = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "You are a Patent Attorney specializing in identify the domain and subject of the patent/Invention and then generate an overall theme",
            },
            {"role": "user", "content": prompt},
        ],
        "max_tokens": 4000,
        "temperature": 0.3,
    }

    response = requests.post(
        f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version={api_version}",
        headers=headers,
        data=json.dumps(data),
    )

    if response.status_code == 200:
        return response.json()["choices"][0]["message"]["content"]
    else:
        logging.warning("Error generating theme!")
        return "Error generating theme."


def extract_text_from_pdf(pdf_file):
    logging.info("Function to extract text from pdf using Fitz (PyMuPDF)")
    pdf_document = fitz.open(pdf_file)
    text_content = []

    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        page_text = page.get_text("text")  # Extracts text from the page
        text_content.append(
            {
                "slide_number": page_number + 1,  # Page numbers start from 1
                "slide_title": f"Page {page_number + 1}",
                "text": page_text.strip(),  # Strips leading/trailing whitespace
            }
        )

    pdf_document.close()
    return text_content


def sanitize_text(text):
    if text:
        sanitized = "".join(
            c
            for c in text
            if c.isprintable()
            and c
            not in {
                "\x00",
                "\x01",
                "\x02",
                "\x03",
                "\x04",
                "\x05",
                "\x06",
                "\x07",
                "\x08",
                "\x0B",
                "\x0C",
                "\x0E",
                "\x0F",
                "\x10",
                "\x11",
                "\x12",
                "\x13",
                "\x14",
                "\x15",
                "\x16",
                "\x17",
                "\x18",
                "\x19",
                "\x1A",
                "\x1B",
                "\x1C",
                "\x1D",
                "\x1E",
                "\x1F",
            }
        )
        return sanitized
    return text


def ensure_proper_spacing(text):
    if text:
        # Fix spacing issues after periods
        text = re.sub(r"\.(?!\s)", ". ", text)  # Ensure space after period
        text = re.sub(r"\s+", " ", text)  # Ensure single space between words
        text = re.sub(
            r"(\.\s+)(\w)", lambda match: match.group(1) + match.group(2).upper(), text
        )  # Capitalize first letter after period
        text = (
            text[0].upper() + text[1:]
        )  # Ensure the first letter of the text is capitalized
    return text


def boldify_text(paragraph, text):
    """Helper function to add text to a paragraph with bold formatting for sections surrounded by '**'."""
    while "**" in text:
        before, _, remaining_text = text.partition("**")
        bold_part, _, after = remaining_text.partition("**")
        paragraph.add_run(before)  # Add regular text before bold part
        bold_run = paragraph.add_run(bold_part)  # Bold text
        bold_run.bold = True
        text = after  # Continue with the rest of the text
    paragraph.add_run(text)  # Add any remaining regular text after the last bold part


def format_content(text):
    """Helper function to remove '###' at the beginning for headings."""
    return text.replace("### ", "").strip()


def save_content_to_word(aggregated_content, output_file_name, extracted_images, theme):
    logging.info("Function to save content into Word Document")
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10.5)  # Reduced font size for paragraphs
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.alignment = 3  # Justify

    doc.add_heading("Overall Theme", level=1).alignment = 0
    theme_paragraph = doc.add_paragraph()
    theme_paragraph.alignment = 0  # Left alignment for paragraph

    # Process and add theme content with bold formatting
    processed_theme = format_content(theme)
    boldify_text(theme_paragraph, processed_theme)

    # # Add the theme at the top of the document
    # doc.add_heading("Overall Theme", level=1).alignment = 0
    # theme_paragraph = doc.add_paragraph(theme)
    # theme_paragraph.alignment = 0  # 0 for Left alignment

    for slide in aggregated_content:
        # Ensure slide is a dictionary
        if isinstance(slide, dict):
            slide_number = slide.get("slide_number", "Unknown Slide")
            slide_title = slide.get("slide_title", "Unknown Title")
            sanitized_title = sanitize_text(slide_title)
            # sanitized_content = sanitize_text(slide.get("insight", ""))
            # properly_spaced_content = ensure_proper_spacing(sanitized_content)

            # Check if slide_number is string or integer
            slide_numbers = (
                slide_number if isinstance(slide_number, str) else f"{slide_number}"
            )

            # doc.add_heading(f"[[{slide_numbers}, {sanitized_title}]]")
            # doc.add_paragraph(f"{properly_spaced_content}")

            # Adding content to the document
            doc.add_heading(f"[[{slide_numbers}, {sanitized_title}]]", level=1)
            doc.add_paragraph(f"{slide['insight']}")
        else:
            st.error(f"Invalid slide structure: {slide}")
            # doc.add_heading(f"[[{slide['slide_numbers']}, {slide['sanitized_title']}]]")
            # doc.add_paragraph(f"{slide['insight']}")

    # Add extracted images after the generated content
    if extracted_images:
        doc.add_heading("Extracted Images", level=1)
        for idx, (image, slide_number) in enumerate(extracted_images):
            _, buffer = cv2.imencode(".png", image)
            image_stream = BytesIO(buffer)
            doc.add_paragraph(f"Image from Slide {slide_number}:")
            doc.add_picture(
                image_stream,
                width=doc.sections[0].page_width
                - doc.sections[0].left_margin
                - doc.sections[0].right_margin,
            )
            doc.add_paragraph("\n")  # Add space after image

    # # Add the theme at the end of the document
    # doc.add_heading("Overall Theme", level=1)
    # doc.add_paragraph(theme)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def extract_and_clean_page_image(page, top_mask, bottom_mask, left_mask, right_mask):
    # Get the page as an image
    logging.info("Function to extract and clean page image")
    pix = page.get_pixmap()
    img_data = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
        pix.height, pix.width, pix.n
    )

    # Convert the image to BGR format for OpenCV
    img_bgr = cv2.cvtColor(img_data, cv2.COLOR_RGB2BGR)

    # Convert to grayscale for processing
    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)

    # Threshold the image to get binary image
    _, binary = cv2.threshold(gray, 240, 255, cv2.THRESH_BINARY_INV)

    # Detect contours of possible images/diagrams
    contours, _ = cv2.findContours(binary, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    # Check if there are any valid contours (image regions)
    valid_contours = [
        cv2.boundingRect(contour)
        for contour in contours
        if cv2.boundingRect(contour)[2] > 50 and cv2.boundingRect(contour)[3] > 50
    ]
    if not valid_contours:
        return None  # Skip the page if no valid images/diagrams are found

    # Create a mask for the detected contours
    mask = np.zeros_like(binary)
    for x, y, w, h in valid_contours:
        # Apply the adjustable top, bottom, left, and right masking values from the sliders
        # Ensure coordinates do not go out of image bounds
        x1 = max(0, x - left_mask)
        y1 = max(0, y - top_mask)
        x2 = min(img_bgr.shape[1], x + w + right_mask)
        y2 = min(img_bgr.shape[0], y + h + bottom_mask)
        cv2.rectangle(mask, (x1, y1), (x2, y2), 255, -1)

    # Use the mask to keep only the regions with images/diagrams
    text_removed = cv2.bitwise_and(img_bgr, img_bgr, mask=mask)

    # Set the background to white where the mask is not applied
    background = np.ones_like(img_bgr) * 255
    cleaned_image = np.where(mask[:, :, None] == 255, text_removed, background)

    # Convert cleaned image to grayscale
    cleaned_image_gray = cv2.cvtColor(cleaned_image, cv2.COLOR_BGR2GRAY)
    return cleaned_image_gray


def extract_images_from_pdf(
    pdf_file, top_mask, bottom_mask, left_mask, right_mask, low_quality_slides
):
    logging.info("Function to extract from PDF")
    # Open the PDF file
    pdf_document = fitz.open(pdf_file)
    page_images = []

    for page_num in range(len(pdf_document)):
        if page_num + 1 in low_quality_slides:
            continue  # Skip low-quality slides

        page = pdf_document.load_page(page_num)

        # Extract and clean the page image
        cleaned_image = extract_and_clean_page_image(
            page, top_mask, bottom_mask, left_mask, right_mask
        )
        if cleaned_image is not None:
            page_images.append(
                (cleaned_image, page_num + 1)
            )  # Keep track of the slide number

    pdf_document.close()
    return page_images


def aggregate_content(text_insights, image_insights, slide_data, continue_insights):
    logging.info("Function to aggregate generated content")
    aggregated_content = []
    processed_slide_numbers = set()

    # Step 1: Add continue insights first, if present
    for continue_slide in continue_insights:
        slide_numbers = continue_slide["slide_number"]  # This may be comma-separated
        slide_title = continue_slide["slide_title"]
        image_insight = continue_slide["insight"]

        # Collect content with slide number and title in separate fields
        aggregated_content.append(
            {
                "slide_number": slide_numbers,  # Store as string if comma-separated
                "slide_title": slide_title,
                "insight": image_insight,
            }
        )

        # Add all slide numbers to processed list, handling comma-separated values
        for slide_num in map(int, slide_numbers.split(",")):
            processed_slide_numbers.add(slide_num)

    # Step 2: Add image insights for slides not processed by continue_insights
    for img in image_insights:
        slide_number = int(img["slide_number"])
        slide_title = img["slide_title"]
        image_insight = img["insight"]

        # Only add if this slide number wasn't already processed by continue_insights
        if slide_number not in processed_slide_numbers:
            aggregated_content.append(
                {
                    "slide_number": str(
                        slide_number
                    ),  # Store as string for consistency
                    "slide_title": slide_title,
                    "insight": image_insight,
                }
            )
            processed_slide_numbers.add(slide_number)

    # Step 3: Add text insights for any remaining unprocessed slides
    for text in text_insights:
        slide_number = int(text["slide_number"])
        slide_title = text["slide_title"]
        text_insight = text["insight"]

        # Only add text insights if the slide number wasn't already processed
        if slide_number not in processed_slide_numbers:
            aggregated_content.append(
                {
                    "slide_number": str(
                        slide_number
                    ),  # Store as string for consistency
                    "slide_title": slide_title,
                    "insight": text_insight,
                }
            )
            processed_slide_numbers.add(slide_number)

    # Step 4: Sort the aggregated content by slide number in ascending order
    # Convert slide_number to tuple of ints for consistent sorting
    aggregated_content = sorted(
        aggregated_content, key=lambda x: tuple(map(int, x["slide_number"].split(",")))
    )

    # Step 5: Return the aggregated content with the required structure
    return aggregated_content


def identify_low_quality_slides(text_content, image_slides):
    logging.info("Identifying low-quality slides.")
    low_quality_slides = set()

    # Ensure all elements in image_slides have a valid 'slide_number' key
    image_slides_dict = {
        int(slide["slide_number"]): slide
        for slide in image_slides
        if "slide_number" in slide
    }  # Extract slide numbers and create a dict for easier access

    for slide in text_content:
        slide_number = int(
            slide["slide_number"]
        )  # Convert slide_number to an integer if it's not already

        # Skip if the slide number is in image_slides
        if slide_number in image_slides_dict:
            # st.write(f"Processing image slide {slide_number}")
            image_slide_data = image_slides_dict[
                slide_number
            ]  # Access the correct slide data

            # Pass the individual slide dictionary to the is_low_quality_image_slide function
            if is_low_quality_image_slide(image_slide_data):
                low_quality_slides.add(slide_number)

            continue  # Skip further checks for image slides

        # Check word count for text slides
        word_count = len(slide["text"].split())
        if word_count < 30:
            low_quality_slides.add(slide_number)

        # Check for generic terms
        if any(
            generic in slide["text"].lower()
            for generic in ["introduction", "thank you", "inventor details", "contents"]
        ):
            low_quality_slides.add(slide_number)

    logging.info(f"Identified {len(low_quality_slides)} low-quality slides.")
    return low_quality_slides


def is_low_quality_image_slide(image_data):
    """Send image slide data to an LLM to check whether the slide is low quality."""
    logging.info("Identifying low-quality slides using LLM")

    base64_image = encode_image(
        image_data["image"]
    )  # Access the image from the individual slide dictionary

    # Create prompt to assess the slide based on the slide number
    prompt = """
    State whether the slide is low quality. 
    If it mostly contains text, check the slide has more than 30 words if it has then consider it high quality, 
    If the slide has less then 20 words without any image in it then consider it low quality.
    If it contains diagrams, figures, graphs, tables, charts, or images, consider it high quality. 
    If the title contains 'summary' then check for whether the slide has more than 30 words if it has then consider it high quality.
    If the title contains 'Inventors', 'introduction', 'contents', 'thank you', or 'inventor details', consider it low quality.
    """

    data = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "You are an image slide quality assessment model.",
            },
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{base64_image}"},
                    },
                ],
            },
        ],
        "max_tokens": 100,
        "temperature": 0.3,
    }

    headers = {
        "Content-Type": "application/json",
        "api-key": api_key,
        "Cache-Control": "no-cache",
        "Pragma": "no-cache",
    }

    response = requests.post(
        f"{azure_endpoint}/openai/deployments/{model}/chat/completions?api-version={api_version}",
        headers=headers,
        data=json.dumps(data),
    )

    if response.status_code == 200:
        assessment = response.json()["choices"][0]["message"]["content"]
        # st.success("low quality" in assessment.lower())
        return "low quality" in assessment.lower()
    else:
        logging.error(
            f"Error processing slide: {response.status_code} to identify whether it is low quality or not"
        )
        st.write(f"Error processing slide: {response.status_code}")
        return False


def upload_to_blob_storage(file_name, file_data):
    try:
        blob_client = blob_service_client.get_blob_client(
            container=container_name, blob=file_name
        )
        blob_client.upload_blob(file_data, overwrite=True)
        # st.info(f"{file_name} uploaded to Azure Blob Storage.")
    except Exception as e:
        st.error(f"Failed to upload {file_name} to Azure Blob Storage: {e}")


def download_from_blob_storage(file_name):
    try:
        blob_client = blob_service_client.get_blob_client(
            container=container_name, blob=file_name
        )
        blob_data = blob_client.download_blob().readall()
        logging.info(f"{file_name} download to Azure Blob Storage.")
        return BytesIO(blob_data)
    except Exception as e:
        st.error(f"Failed to download {file_name} from Azure Blob Storage: {e}")
        return None


# Replace disallowed words
def replace_disallowed_words(text):
    disallowed_words = {
        "necessary": "required",
        "necessity": "requirement",
        "necessitate": "require",
        "necessitated": "required",
        "necessitates": "requires",
        "necessarily": "inevitably",
        "necessitating": "requiring",
        "contain": "include",
        "critical": "captious",
    }
    for word, replacement in disallowed_words.items():
        text = text.replace(word, replacement)
    # Ensure single paragraph output
    text = " ".join(text.split())

    return text


# Streamlit app interface update
def main():
    st.title("PATENT APPLICATION")

    if "top_mask" not in st.session_state:
        st.session_state.top_mask = 40
    if "bottom_mask" not in st.session_state:
        st.session_state.bottom_mask = 40
    if "left_mask" not in st.session_state:
        st.session_state.left_mask = 85
    if "right_mask" not in st.session_state:
        st.session_state.right_mask = 85

    st.markdown(
        """
        <style>
        .stButton>button {
            padding: 10px 20px;
            width: 140px;
            height: 40px;
            font-size: 16px;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    col1, col2 = st.sidebar.columns(2)
    with col1:
        if st.button("Default"):
            st.session_state.top_mask = 40
            st.session_state.bottom_mask = 40
            st.session_state.left_mask = 85
            st.session_state.right_mask = 85

    with col2:
        if st.button("A4"):
            st.session_state.top_mask = 70
            st.session_state.bottom_mask = 70
            st.session_state.left_mask = 85
            st.session_state.right_mask = 85

    top_mask = st.sidebar.slider(
        "Adjust Top Masking Value",
        min_value=10,
        max_value=100,
        value=st.session_state.top_mask,
        step=1,
    )
    bottom_mask = st.sidebar.slider(
        "Adjust Bottom Masking Value",
        min_value=10,
        max_value=100,
        value=st.session_state.bottom_mask,
        step=1,
    )
    left_mask = st.sidebar.slider(
        "Adjust Left Masking Value",
        min_value=10,
        max_value=500,
        value=st.session_state.left_mask,
        step=1,
    )
    right_mask = st.sidebar.slider(
        "Adjust Right Masking Value",
        min_value=10,
        max_value=200,
        value=st.session_state.right_mask,
        step=1,
    )

    if (
        top_mask != st.session_state.top_mask
        or bottom_mask != st.session_state.bottom_mask
        or left_mask != st.session_state.left_mask
        or right_mask != st.session_state.right_mask
    ):
        st.session_state.top_mask = top_mask
        st.session_state.bottom_mask = bottom_mask
        st.session_state.left_mask = left_mask
        st.session_state.right_mask = right_mask

    # File uploader for PDF

    uploaded_ppt = st.file_uploader("Upload a PPT file", type=["pptx"])

    text_length = st.select_slider(
        "Select creativity level",
        options=["Standard", "Blend", "Creative"],
        value="Blend",
    )

    if st.button("Start Generate"):
        # Extract the base name of the uploaded PPT file
        if uploaded_ppt is None:
            st.error("Please upload a PPT file before proceeding.")
            return

        ppt_filename = uploaded_ppt.name
        base_filename = os.path.splitext(ppt_filename)[0]
        output_word_filename = f"{base_filename}.docx"

        # Save uploaded PPT to a temporary file and upload to Azure Blob Storage
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pptx") as temp_ppt_file:
            temp_ppt_file.write(uploaded_ppt.read())
            temp_ppt_file_path = temp_ppt_file.name

        upload_to_blob_storage(ppt_filename, open(temp_ppt_file_path, "rb"))

        # Download the PPT file from Azure Blob Storage
        ppt_blob = download_from_blob_storage(ppt_filename)
        if not ppt_blob:
            st.error("Failed to download PPT file from Azure Blob Storage.")
            return

        # Convert PPT to PDF
        with open("uploaded_ppt.pptx", "wb") as f:
            f.write(ppt_blob.read())

        if not ppt_to_pdf("uploaded_ppt.pptx", "uploaded_pdf.pdf"):
            st.error("PDF conversion failed. Please check the uploaded PPT file.")
            return

        text_content = extract_text_from_pdf("uploaded_pdf.pdf")
        # Extract images

        title_slide_images = extract_slide_images_for_title_extraction(
            "uploaded_pdf.pdf"
        )
        image_content = detect_images_from_pdf("uploaded_pdf.pdf")

        low_quality_slides = identify_low_quality_slides(text_content, image_content)
        # st.write(low_quality_slides)

        slide_data = extract_titles_from_images(title_slide_images)

        continued_check = continued_title_check(slide_data)
        if continued_check:
            for entry in continued_check:
                slide_sets = entry["set_of_slides"].strip("[]").split("], [")
                st.sidebar.write("Continued Slide Sets:")
                for slide_set in slide_sets:
                    st.sidebar.write(f"[{slide_set.strip()}]")

        if image_content:
            # Convert low-quality slides input into list
            low_quality_slides = [
                int(slide) for slide in low_quality_slides if isinstance(slide, int)
            ]

            # Step 3: Continue with generating insights or further processing using the slide_data
            combined_text = extract_text_and_titles_from_pdf("uploaded_pdf.pdf")
            overall_theme = generate_overall_theme(combined_text)

            system_prompt = generate_prompt(overall_theme)
            # Generate insights via LLM

            text_insights = generate_text_insights(
                text_content,
                text_length,
                low_quality_slides,
                slide_data,
                system_prompt,
            )
            insights = generate_image_insights(
                image_content,
                text_length,
                low_quality_slides,
                system_prompt,
                slide_data,
            )

            continue_insights = []
            if continued_check != []:
                # st.warning("Inside")
                continue_insights = generate_continue_insights(
                    text_length,
                    continued_check,
                    image_content,
                    text_content,
                    slide_data,
                    system_prompt,
                )

            st.write(
                "--------------------------------------------------------------------------------------------------------------------------------"
            )

            extracted_images = extract_images_from_pdf(
                "uploaded_pdf.pdf",
                top_mask,
                bottom_mask,
                left_mask,
                right_mask,
                low_quality_slides,
            )

            aggregated_content = aggregate_content(
                text_insights, insights, slide_data, continue_insights
            )

            for insight in aggregated_content:
                st.subheader(f"[[{insight['slide_number']}, {insight['slide_title']}]]")
                st.markdown(insight["insight"])

            st.info("Saving to Word document...")
            output_doc = save_content_to_word(
                aggregated_content,
                output_word_filename,
                extracted_images,
                overall_theme,
            )

            st.download_button(
                label="Download Word Document",
                data=output_doc,
                file_name=output_word_filename,
            )
            st.success("Processing completed successfully!")
        else:
            st.warning("No images, flowcharts, or diagrams detected in the PDF.")


if __name__ == "__main__":
    main()
